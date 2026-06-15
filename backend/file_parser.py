"""
file_parser.py
Extracts plain text from PDF, DOC, and DOCX files.
"""

import os
import subprocess
import tempfile
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_uploaded_file(file_path: str) -> str:
    """
    Detect file type and extract text content.
    Returns extracted text as a string.
    """
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    elif ext == ".doc":
        return _extract_doc(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(file_path: str) -> str:
    """Extract text from PDF using pdfminer."""
    try:
        from pdfminer.high_level import extract_text
        text = extract_text(file_path)
        return text.strip()
    except ImportError:
        logger.warning("pdfminer not available, trying pypdf fallback")
        return _extract_pdf_pypdf(file_path)
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise ValueError(f"Failed to read PDF: {e}")


def _extract_pdf_pypdf(file_path: str) -> str:
    """Fallback PDF extraction using pypdf."""
    try:
        import pypdf
        reader = pypdf.PdfReader(file_path)
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        return "\n".join(pages).strip()
    except Exception as e:
        raise ValueError(f"Failed to read PDF with fallback: {e}")


def _extract_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        import docx
        doc = docx.Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError(f"Failed to read DOCX: {e}")


def _extract_doc(file_path: str) -> str:
    """Extract text from legacy .doc using antiword or fallback."""
    # Try antiword (Linux)
    try:
        result = subprocess.run(
            ["antiword", file_path],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Try converting to docx via python-docx2txt
    try:
        import docx2txt
        text = docx2txt.process(file_path)
        if text and text.strip():
            return text.strip()
    except Exception:
        pass

    raise ValueError(
        "Could not read .doc file. Please convert it to .docx or PDF and try again."
    )
