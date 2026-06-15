"""
docx_parser.py — Deka Research Questionnaire Parser v10 (Generic)
==================================================================
General — tidak hardcode topik/merek/industri.

Perubahan v10:
- _normalize_route: handle multi-value route (OE SA, MA SA, SA MA, dll)
- _infer_route_from_text: inferensi tipe dari pola bahasa Indonesia universal
- _peek_choices_from_cell: peek choices sebelum inferensi route agar urutan benar
- Slogan regex Qxxb fix: support dua+ digit (Q16b, Q17b, dst)
- Slogan Qxxa (SA ya/tidak) dan Qxxc (rating skala) otomatis terdeteksi
- note type untuk instruksi DP murni ("tidak perlu ditanyakan", "interviewer perhatikan")
- Tempat beli (V6/Z6/Q6 tanpa sub) → MA otomatis dari teks
- Alasan rekomendasi / TIDAK rekomendasi → OE dari teks
- P-section fallback cerdas: ada choices → SA/MA, tidak ada → OE/text
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

import docx
from docx.table import Table, _Cell

# ── Pattern & konstanta ────────────────────────────────────────────────────────

_Q_ID_RE = re.compile(
    r'^('
    r'SEC\d+'
    r'|S\d+[a-zA-Z]?'
    r'|MD\d*[a-zA-Z0-9\-\.]*'
    r'|[A-Z]{1,2}\d+[a-zA-Z0-9\-\.]*'
    r'|[A-Z]{1,2}\d*\.[a-zA-Z0-9]?'
    r'|[VYZF]\d*[a-zA-Z0-9\-\.]*'
    r'|[A-Z]\d*'
    r')$',
    re.IGNORECASE
)

# Marker instruksi → hint bukan label
_HINT_MARKERS = [
    "KARTU BANTU", "TUNJUKKAN", "SPONTAN", "PROBE", "BACAKAN",
    "ROTASIKAN", "INTERVIEWER", "DP:", "NOTE:", "NOTE TO DP",
    "TANYAKAN KEPADA", "TANYAKAN JIKA", "DITANYAKAN JIKA",
    "CATAT USIA", "CATAT AKTUAL", "TRANSFER JAWABAN",
    "JAWABAN SPONTAN",
]

_STOP_WORDS     = {"AKHIRI WAWANCARA", "HENTIKAN WAWANCARA", "STOP", "STOP WAWANCARA"}
_CONTINUE_WORDS = {"LANJUTKAN", "CONTINUE", "CEK KUOTA"}

_NON_SECTION_RE = re.compile(
    r'^(TABEL ISIAN|PANEL KONTROL|KARTU BANTU P|NOTE:|CATATAN DP|DP:|'
    r'\[.*\]|AKHIR DARI|MASUKAN LIST)',
    re.IGNORECASE
)

# ── Brand & product choices — dinamis, di-discover dari dokumen ───────────────
#
# Struktur baru: satu dict global yang menyimpan semua merek per "slot produk".
# Slot A1 = produk utama, A2 = produk kedua (jika ada), Y1/V/Z = produk ketiga, dst.
# Diisi oleh _extract_brand_lists_from_doc() saat parse dimulai.
#
# Fallback: daftar generik yang dipakai jika brand grid tidak ditemukan di dokumen.

_GENERIC_BRAND_FALLBACK = [
    ("1", "Merek A"), ("2", "Merek B"), ("3", "Merek C"), ("4", "Merek D"),
    ("5", "Merek E"), ("98", "Lainnya"), ("99", "Tidak tahu"),
]

# Slot runtime — diisi oleh _extract_brand_lists_from_doc()
_BRAND_SLOTS: dict[str, list[tuple]] = {
    "primary":   list(_GENERIC_BRAND_FALLBACK),
    "secondary": list(_GENERIC_BRAND_FALLBACK),
    "tertiary":  list(_GENERIC_BRAND_FALLBACK),
}
_PRODUCT_LABELS: dict[str, str] = {
    "primary":   "Produk Utama",
    "secondary": "Produk Kedua",
    "tertiary":  "Produk Ketiga",
}
# Choices produk kategori (S8/PHOTOCARD) — diisi dari dokumen atau generic
_CHOICES_PRODUK: list[tuple] = [
    ("1", "Produk Utama"), ("2", "Produk Kedua"), ("3", "Produk Ketiga"),
]
_CHOICES_JENIS_PRODUK: list[tuple] = [
    ("1", "Tipe A"), ("2", "Tipe B"), ("3", "Tipe C"),
    ("98", "Lainnya"), ("99", "Tidak Tahu"),
]

# Alias backward-compat (dipakai di banyak tempat di parser lama)
def _get_merek(slot: str = "primary") -> list[tuple]:
    return _BRAND_SLOTS.get(slot, _GENERIC_BRAND_FALLBACK)

# Skala 1-10
_SKALA_10 = [(str(i), str(i)) for i in range(1, 11)]
# Skala 0-10 (NPS)
_SKALA_11 = [(str(i), str(i)) for i in range(0, 11)]

# Usia range
_USIA_RANGE = [
    ("1","Di bawah 25 tahun"), ("2","25–29 tahun"), ("3","30–34 tahun"),
    ("4","35–39 tahun"), ("5","40–44 tahun"), ("6","45–50 tahun"),
    ("7","Di atas 50 tahun"),
]

# Time range sejak ganti merek
_TIME_RANGE = [
    ("1","Di bawah 6 bulan yang lalu"), ("2","Di bawah 1 tahun yang lalu"),
    ("3","Di bawah 2 tahun yang lalu"), ("4","Di bawah 3 tahun yang lalu"),
    ("5","Lebih dari 3 tahun yang lalu"),
]


# ── Helpers ────────────────────────────────────────────────────────────────────

def _cell_text(cell: _Cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs).strip()


def _detect_routing(text: str) -> str:
    t = text.strip().upper()
    for w in _STOP_WORDS:
        if w in t:
            return "STOP"
    for w in _CONTINUE_WORDS:
        if w in t:
            return "CONTINUE"
    return t if t else ""


def _normalize_route(raw: str) -> str:
    """
    Normalisasi kolom route dari dokumen.

    Strategi multi-value (mis. 'OE SA', 'MA SA', 'SA MA'):
    - Kalau ada SA → SA menang (pertanyaan tunggal/terpilih)
    - Kalau ada OE dan tidak ada SA/MA → OE
    - Kalau ada MA tapi tidak SA → MA
    Ini mencerminkan konvensi Deka: kalau satu sel punya dua mode,
    mode paling restriktif (SA) biasanya yang dimaksud untuk XLSForm.
    Pengecualian: 'MA OE' → tetap MA+OE (ada sub open-ended).
    """
    raw = raw.strip().upper().replace("\n", " ").replace("  ", " ")
    raw = re.sub(r'\s+', ' ', raw).strip()

    explicit_map = {
        "SA": "SA", "MA": "MA", "OE": "OE",
        "M": "M (OE)", "M (OE)": "M (OE)",
        "SA MA": "SA", "MA SA": "SA",
        "SA OE": "SA", "OE SA": "OE",
        "MA OE": "MA+OE", "OE MA": "MA+OE",
        "SA MA MA": "SA", "SA SA": "SA",
        "MA MA": "MA",
    }
    if raw in explicit_map:
        return explicit_map[raw]

    # Fallback: multi-token, ambil token pertama yang dikenali
    tokens = raw.split()
    priority = ["SA", "MA", "OE"]
    for p in priority:
        if p in tokens:
            return p
    return raw


# ── Pola inferensi tipe dari teks (general, bahasa Indonesia) ─────────────────
#
# Tiap entry: (compiled_regex, route_type, keterangan)
# Urutan PENTING — pertama yang match menang.
#
_TEXT_INFERENCE_RULES: list[tuple[re.Pattern, str]] = [
    # Instruksi DP murni → note (bukan pertanyaan)
    (re.compile(
        r'^(tidak perlu ditanyakan|interviewer\s+perhatikan|catatan\s+dp'
        r'|tabel\s+isian|panel\s+kontrol|transfer\s+jawaban)',
        re.I), "note"),

    # Pertanyaan rating/skala → SA (akan ditangani _parse_rating_q via auto-detect)
    (re.compile(r'seberapa\s+(sesuai|yakin|puas|setuju|sering|penting|baik)'
                r'.*skala|dengan\s+skala\s+\d', re.I | re.S), "SA_RATING"),

    # Open-ended: alasan, mengapa, apa alasan
    (re.compile(r'^(mengapa\s+(anda\s+)?(tidak\s+)?|apa\s+alasan\s+|'
                r'apa\s+yang\s+membuat\s+|jelaskan\s+|ceritakan\s+)', re.I), "OE"),

    # Open-ended: berapa persen / persentase → integer/OE
    (re.compile(r'berapa\s+(persen|%|persentase)', re.I), "OE"),

    # Open-ended: berapa lama, berapa rata-rata (angka bebas)
    (re.compile(r'^berapa\s+(lama|rata.rata|jumlah|total|nilai)', re.I), "OE"),

    # SA: apakah pernah / apakah anda (ya/tidak)
    (re.compile(r'^apakah\s+(anda\s+)?(pernah|masih|sudah|akan|ada|mau|bisa|'
                r'toko\s+anda|perusahaan\s+anda)', re.I), "SA"),

    # SA: manakah dari ... yang PALING (satu jawaban)
    (re.compile(r'manakah\s+dari.*yang\s+paling\b|yang\s+paling\s+(sering|banyak|'
                r'sesuai|disukai|direkomendasikan|anda\s+fav)', re.I | re.S), "SA"),

    # SA: merek apa yang PALING (satu)
    (re.compile(r'merek\s+\w[\w\s/]*\s+apa\s+yang\s+paling\b|'
                r'paling\s+anda\s+rekomendasikan', re.I), "SA"),

    # SA: untuk merek apa slogan tersebut
    (re.compile(r'untuk\s+merek\s+apa\s+slogan', re.I), "SA"),

    # SA: siapakah yang ... / siapa yang
    (re.compile(r'^siapa(kah)?\s+yang\b', re.I), "SA"),

    # SA: sudah berapa lama / sejak kapan (pilihan range)
    (re.compile(r'^sudah\s+berapa\s+lama\b|^sejak\s+(kapan|usia\s+berapa)\b', re.I), "SA"),

    # MA: di mana saja / tempat apa saja (membeli/mendapatkan)
    (re.compile(r'di\s+mana\s+saja\b|tempat\s+(apa\s+saja|mana\s+saja)\b', re.I), "MA"),

    # MA: promosi apa saja / promosi apa yang
    (re.compile(r'promosi\s+apa\s+(saja|yang)\b', re.I), "MA"),

    # MA: hal-hal apa / hal apa saja
    (re.compile(r'hal.hal\s+apa(kah)?\s+|hal\s+apa\s+saja\b', re.I), "MA"),

    # MA: kegiatan apa saja / jenis apa saja
    (re.compile(r'(kegiatan|jenis|merek|ukuran)\s+\w[\w\s/]*\s+apa\s+saja\b', re.I), "MA"),

    # MA: merek-merek apa / merek apa yang anda ketahui (multi)
    (re.compile(r'merek.merek\s+\w[\w\s/]*\s+(apa|mana)\b|'
                r'merek\s+\w[\w\s/]*\s+apa\s+(yang\s+)?(anda\s+)?(ketahui|kenal|'
                r'pernah\s+gunakan|jual\s+saat\s+ini)\b', re.I), "MA"),

    # MA: mana saja yang TIDAK anda pertimbangkan
    (re.compile(r'yang\s+tidak\s+(anda\s+)?pertimbangkan\b|'
                r'tidak\s+akan\s+(anda\s+)?pertimbangkan\b', re.I), "MA"),

    # SA: default untuk pertanyaan pilihan tunggal yang tidak tertangkap di atas
    # (akan hanya berlaku jika ada choices)
]


def _infer_route_from_text(q_text: str, hint: str,
                           has_choices: bool, choices: list[dict]) -> str:
    """
    Inferensi route_type dari pola teks pertanyaan.
    Dipanggil hanya jika kolom route dokumen kosong.

    Returns route string yang kompatibel dengan _ROUTE_TO_TYPE di json_to_xlsform,
    atau "" jika tidak bisa diinfer (biarkan fallback di json_to_xlsform).
    """
    combined = (q_text + " " + hint).strip()
    if not combined:
        return "SA" if has_choices else ""

    for pattern, route in _TEXT_INFERENCE_RULES:
        if pattern.search(combined):
            if route == "note":
                return "note"
            if route == "SA_RATING":
                return "SA"  # akan auto-detect skala di _parse_question_row
            if route == "OE":
                return "OE"
            if route in ("SA", "MA"):
                return route

    # Tidak ada pola cocok — gunakan keberadaan choices sebagai signal
    if has_choices:
        # Cek teks untuk sinyal multi: "saja", "apalagi", "(M)", semua yang
        multi_signals = re.compile(
            r'\bsaja\b|\bapalagi\b|\(\s*M\s*\)|\bpilih\s+semua\b|'
            r'\bmulti\b|\bboleh\s+lebih\s+dari\s+satu\b', re.I
        )
        if multi_signals.search(combined):
            return "MA"
        return "SA"

    return ""  # biarkan json_to_xlsform handle


def _peek_choices_from_cell(cell: "_Cell") -> list[dict]:
    """
    Intip choices dari nested table di sel tanpa full parse.
    Dipakai untuk inferensi route sebelum parse lengkap.
    Returns list (bisa kosong).
    """
    choices = []
    for nt in cell.tables:
        nc = len(nt.columns)
        if nc >= 10:  # skala — abaikan untuk count
            continue
        for row in nt.rows:
            cells = [c.text.strip() for c in row.cells]
            if nc >= 2:
                c0, c1 = cells[0], cells[1] if len(cells) > 1 else ""
                if c0 and not re.match(r'^(ROTASIKAN|DP:|KARTU|\[)', c0, re.I):
                    choices.append({"code": c1, "label": c0})
            elif nc == 1 and cells[0]:
                if not re.match(r'^(ROTASIKAN|DP:|KARTU|\[)', cells[0], re.I):
                    choices.append({"code": "", "label": cells[0]})
        if choices:
            break
    return choices


def _split_hint(text: str) -> tuple[str, str]:
    """Pisahkan teks pertanyaan dari instruksi interviewer."""
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    q_lines, h_lines = [], []
    for ln in lines:
        upper = ln.upper()
        if any(upper.startswith(m) for m in _HINT_MARKERS):
            h_lines.append(ln)
        elif re.match(r'^\[.*\]$', ln):   # instruksi dalam kurung kotak
            h_lines.append(ln)
        else:
            q_lines.append(ln)
    return " ".join(q_lines).strip(), " ".join(h_lines).strip()


def _make_q(q_id, section, subsection, question, hint="",
            route_type="SA", choices=None, skip_logic="",
            is_grid=False, is_header=False, raw_text="") -> dict:
    return {
        "id":          q_id,
        "section":     section,
        "subsection":  subsection,
        "question":    question,
        "hint":        hint,
        "route_type":  route_type,
        "choices":     choices or [],
        "skip_logic":  skip_logic,
        "is_grid":     is_grid,
        "is_header":   is_header,
        "raw_text":    raw_text,
    }


def _choices_from_list(pairs: list[tuple[str, str]]) -> list[dict]:
    return [{"code": code, "label": label, "routing": ""} for code, label in pairs]


def _parse_simple_nested_table(tbl: Table) -> list[dict]:
    """Parse tabel 2-kolom sederhana → choices."""
    choices = []
    nc = len(tbl.columns)
    for row in tbl.rows:
        cells = [c.text.strip() for c in row.cells]
        if nc == 1:
            if cells[0] and not re.match(r'^\[|^ROTASIKAN|^DP:', cells[0]):
                choices.append({"code": "", "label": cells[0], "routing": ""})
        elif nc >= 2:
            # Coba deteksi label vs code
            c0, c1 = cells[0], cells[1] if len(cells) > 1 else ""
            routing = cells[2] if len(cells) > 2 else ""
            # Skip header rows
            if re.match(r'^ROTASIKAN|^\[DP|^KARTU', c0, re.I):
                continue
            if re.match(r'^\d+$', c1.strip()):
                label, code = c0, c1
            elif re.match(r'^\d+$', c0.strip()):
                code, label = c0, c1
            else:
                label, code = c0, c1
            if label or code:
                choices.append({
                    "code":    code.strip(),
                    "label":   label.strip(),
                    "routing": _detect_routing(routing),
                })
    # Dedup
    seen, deduped = set(), []
    for ch in choices:
        key = (ch["code"], ch["label"])
        if key not in seen and (ch["code"] or ch["label"]):
            seen.add(key)
            deduped.append(ch)
    return deduped


def _parse_brand_grid_table(nt: Table) -> dict[str, list[dict]]:
    """
    Parse TABEL ISIAN brand grid (9 col × N rows).
    Header row 0: ['', 'Q1', 'Q2', 'Q3', 'Q4a', 'Q4b', 'Q5', 'Q6a', 'Q6c'] (contoh)
    Header row 1: ['', 'PERNAH (M)', '1 TAHUN (M)', ...] — deskripsi sub
    Data rows  2+: ['Alderon', '1', '1', '1', ...]

    Returns dict: { 'q1': [{code:'1',label:'Alderon'},...], 'q2': [...], ... }
    Merek-kode di kolom brand otomatis dipetakan per pertanyaan.
    """
    if len(nt.rows) < 3:
        return {}

    rows = list(nt.rows)
    header_ids = [c.text.strip() for c in rows[0].cells]
    # Normalisasi ID kolom: 'Q6a' → 'q6a', 'V6a' → 'v6a', dll
    col_ids = [h.lower().replace("-","_").replace(" ","") for h in header_ids]

    result: dict[str, list[dict]] = {cid: [] for cid in col_ids if cid}

    for data_row in rows[2:]:  # skip 2 baris header
        cells = [c.text.strip() for c in data_row.cells]
        brand_label = cells[0]
        if not brand_label or re.match(r'^TT|^DP:|^\[', brand_label, re.I):
            # TT/TA biasanya baris terakhir, tetap include
            if brand_label.upper().startswith("TT"):
                brand_label = "TT/TA"
            else:
                continue

        for col_idx, cid in enumerate(col_ids):
            if not cid or col_idx >= len(cells):
                continue
            code = cells[col_idx].strip()
            if code and re.match(r'^\d+$', code):
                result.setdefault(cid, []).append({
                    "code":    code,
                    "label":   brand_label,
                    "routing": "",
                })

    # Bersihkan entry kosong
    return {k: v for k, v in result.items() if v}


def _extract_brand_lists_from_doc(doc: docx.Document) -> None:
    """
    Ekstrak brand list dari dokumen secara generik.

    Strategi:
    1. Temukan brand grid (nested table 6+ kolom dengan header Q1/Q2/V1/Z1 dll.)
    2. Slot primary  ← brand grid dengan Q1/Q2 di header
    3. Slot secondary ← brand grid dengan V1/V2 di header (produk kedua)
    4. Slot tertiary  ← brand grid dengan Z1/Z2 di header (produk ketiga)
    5. Jika ada label produk di baris sebelum grid → gunakan sebagai product label
    6. PHOTOCARD (S8c) dan jenis produk (S8d/S8e) → choices dinamis
    """
    global _BRAND_SLOTS, _PRODUCT_LABELS, _CHOICES_PRODUK, _CHOICES_JENIS_PRODUK

    def _extract_col0_pairs(nt: Table, skip_rows: int = 2) -> list[tuple]:
        pairs = []
        for row in list(nt.rows)[skip_rows:]:
            cells = [c.text.strip() for c in row.cells]
            label = cells[0]
            code  = cells[1] if len(cells) > 1 else ""
            if not label:
                continue
            if re.match(r'^(ROTASIKAN|DP:|TT/)', label, re.I) and not code:
                continue
            pairs.append((code, label))
        return [(c, l) for c, l in pairs if c or l]

    def _infer_product_label(before_text: str) -> str:
        """Coba ekstrak nama produk dari teks sebelum brand grid."""
        # Cari pola "merek XXX", "brand XXX", atau nama produk dalam caps
        m = re.search(r'(?:merek|brand|produk)\s+([A-Za-z0-9/ ]{3,40})', before_text, re.I)
        if m:
            return m.group(1).strip().title()
        # Fallback: ambil kata kapital > 3 huruf
        caps = re.findall(r'\b[A-Z][A-Za-z0-9]{2,}\b', before_text)
        if caps:
            return caps[0]
        return ""

    slots_assigned = {"primary": False, "secondary": False, "tertiary": False}

    for tbl in doc.tables:
        rows = list(tbl.rows)
        for i, row in enumerate(rows):
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 2:
                continue
            q_id   = cells[0].upper()
            q_cell = row.cells[1]

            # ── Brand grid dari baris kosong (q_id kosong) ────────────────
            if not cells[0]:
                for nt in q_cell.tables:
                    if len(nt.columns) < 6 or len(nt.rows) < 3:
                        continue
                    col_headers = [c.text.strip().upper() for c in nt.rows[0].cells]
                    pairs = _extract_col0_pairs(nt, skip_rows=2)
                    if not pairs:
                        continue
                    # Teks baris sebelumnya untuk inferensi nama produk
                    prev_text = rows[i-1].cells[1].text if i > 0 else ""

                    if "Q1" in col_headers and not slots_assigned["primary"]:
                        _BRAND_SLOTS["primary"] = pairs
                        lbl = _infer_product_label(prev_text)
                        if lbl: _PRODUCT_LABELS["primary"] = lbl
                        slots_assigned["primary"] = True
                    elif "V1" in col_headers and not slots_assigned["secondary"]:
                        _BRAND_SLOTS["secondary"] = pairs
                        lbl = _infer_product_label(prev_text)
                        if lbl: _PRODUCT_LABELS["secondary"] = lbl
                        slots_assigned["secondary"] = True
                    elif "Z1" in col_headers and not slots_assigned["tertiary"]:
                        _BRAND_SLOTS["tertiary"] = pairs
                        lbl = _infer_product_label(prev_text)
                        if lbl: _PRODUCT_LABELS["tertiary"] = lbl
                        slots_assigned["tertiary"] = True
                    elif not slots_assigned["primary"]:
                        # Grid pertama yang ditemukan tanpa header Q/V/Z jelas
                        _BRAND_SLOTS["primary"] = pairs
                        lbl = _infer_product_label(prev_text)
                        if lbl: _PRODUCT_LABELS["primary"] = lbl
                        slots_assigned["primary"] = True

                    # S8d/S8e (3-kolom: jenis produk)
                    if len(nt.columns) == 3:
                        h = [c.text.strip().upper() for c in nt.rows[0].cells]
                        if "S8D" in h or "S8E" in h:
                            pairs_sde = []
                            for r in list(nt.rows)[1:]:
                                c0 = r.cells[0].text.strip()
                                c1 = r.cells[1].text.strip() if len(r.cells) > 1 else ""
                                if c0 and c1:
                                    pairs_sde.append((c1, c0))
                            if pairs_sde:
                                _CHOICES_JENIS_PRODUK[:] = pairs_sde

            # ── S8c nested table → PHOTOCARD produk ──────────────────────
            if q_id in ("S8C", "S8B", "S8A"):
                for nt in q_cell.tables:
                    if len(nt.columns) < 3 or len(nt.rows) < 3:
                        continue
                    pairs_s8 = []
                    for r in list(nt.rows)[1:]:
                        label = r.cells[0].text.strip()
                        code  = r.cells[1].text.strip() if len(r.cells) > 1 else ""
                        if label and code and re.match(r'^\d+$', code):
                            pairs_s8.append((code, label))
                    if pairs_s8:
                        _CHOICES_PRODUK[:] = pairs_s8

    # Jika primary brand berhasil diisi, update secondary/tertiary fallback juga
    # agar tidak pakai "Merek A/B/C" kalau ada satu brand list saja
    if slots_assigned["primary"] and not slots_assigned["secondary"]:
        _BRAND_SLOTS["secondary"] = _BRAND_SLOTS["primary"]
        _BRAND_SLOTS["tertiary"]  = _BRAND_SLOTS["primary"]



def _is_screening_table(tbl: Table) -> bool:
    """
    Deteksi tabel screening 3-kolom.
    v10 fix: lebih lenient — header kolom pertama tidak harus persis "NO",
    bisa kosong atau berisi angka/teks lain (seperti di kuesioner Tukang).
    """
    if len(tbl.columns) != 3 or len(tbl.rows) < 2:
        return False
    h = [c.text.strip().upper() for c in tbl.rows[0].cells]
    # Kondisi 1 (strict): header klasik Deka
    if h[0] in ("NO", "NO.") and "PERTANYAAN" in h[1] and "ROUTE" in h[2]:
        return True
    # Kondisi 2 (lenient): header kolom 1 mengandung "PERTANYAAN" dan kolom 2 "ROUTE"
    if "PERTANYAAN" in h[1] and "ROUTE" in h[2]:
        return True
    # Kondisi 3 (fallback): row pertama atau kedua punya q_id valid di kolom 0
    for row in list(tbl.rows)[:3]:
        r = [c.text.strip() for c in row.cells]
        if r[0] and _Q_ID_RE.match(r[0]):
            return True
    return False


def _is_main_table(tbl: Table) -> bool:
    if len(tbl.columns) != 3 or len(tbl.rows) < 3:
        return False
    for row in tbl.rows[:5]:
        cells = [c.text.strip() for c in row.cells]
        if "KUESIONER UTAMA" in cells[1].upper():
            return True
    return False


def _is_info_table(tbl: Table) -> bool:
    """Deteksi tabel Informasi Responden dari halaman pertama kuesioner Deka."""
    if len(tbl.rows) < 2:
        return False
    # Kumpulkan semua teks dari 8 baris pertama
    all_text = ""
    for row in list(tbl.rows)[:8]:
        for cell in row.cells:
            all_text += " " + cell.text.strip().upper()
    # Tabel info responden biasanya berisi kata-kata ini
    info_markers = ["NAMA RESPONDEN", "INTERVIEWER", "TANGGAL WAWANCARA",
                    "NOMOR KUESIONER", "WAKTU MULAI", "NOMOR RESPONDEN"]
    return sum(1 for m in info_markers if m in all_text) >= 2


def _extract_kota_from_doc(doc: docx.Document) -> list[tuple[str, str]]:
    """
    Ekstrak daftar kota dari pertanyaan screening kota di dokumen (S6 atau mirip).

    Format dokumen Deka Research: "NamaKota<spasi>Kode" di kanan (mis. "Jakarta 1")
    BUKAN "1. NamaKota" — itulah bug v8 yang diperbaiki di v10.

    Returns list of (code, label).
    """
    KOTA_PATTERNS = [
        r'jabodetabek', r'jakarta', r'bogor', r'depok', r'tangerang', r'bekasi',
        r'bandung', r'cirebon', r'semarang', r'tegal', r'solo', r'surakarta',
        r'surabaya', r'malang', r'yogyakarta', r'madiun',
        r'medan', r'palembang', r'makassar', r'denpasar', r'balikpapan',
        r'samarinda', r'pontianak', r'manado', r'pekanbaru', r'batam',
        r'bandar lampung', r'padang',
    ]
    kota_re = re.compile('|'.join(KOTA_PATTERNS), re.IGNORECASE)

    found: list[tuple[str, str]] = []
    seen_labels: set[str] = set()

    def _parse_cell(txt: str) -> list[tuple[str, str]]:
        results = []
        seen = set()
        lines = [l.strip() for l in re.split(r'[\n\t]', txt) if l.strip()]
        for ln in lines:
            # Format A (Deka standard): "NamaKota ... Kode" — nama di kiri, angka di kanan
            m_a = re.match(r'^([A-Za-z][A-Za-z\s/\-]{1,30}?)\s+(\d{1,2})\s*$', ln)
            if m_a:
                label, code = m_a.group(1).strip(), m_a.group(2).strip()
                if kota_re.search(label) and label.upper() not in seen:
                    seen.add(label.upper())
                    results.append((code, label))
                continue
            # Format B (alternatif): "Kode. NamaKota"
            m_b = re.match(r'^(\d{1,2})[.\s]\s*([A-Za-z][A-Za-z\s/\-]{2,30})\s*$', ln)
            if m_b:
                code, label = m_b.group(1).strip(), m_b.group(2).strip()
                if kota_re.search(label) and label.upper() not in seen:
                    seen.add(label.upper())
                    results.append((code, label))
                continue
            # Format C: satu baris panjang "Jakarta 1 Bogor 2 Depok 3 ..."
            if kota_re.search(ln) and len(ln) > 15:
                for label, code in re.findall(
                    r'([A-Z][a-zA-Z\s/\-]{2,20}?)\s+(\d{1,2})(?=\s+[A-Z]|\s*$|\s+\*\*)',
                    ln
                ):
                    label = label.strip()
                    if kota_re.search(label) and label.upper() not in seen:
                        seen.add(label.upper())
                        results.append((code, label))
        return results

    # Pass 1: cari baris S6 / S0b / "kota tempat tinggal" di tabel screening
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            q_id_cell = cells[0].upper() if cells else ""
            is_kota_row = (
                re.match(r'^S0?6$', q_id_cell)
                or re.match(r'^S0[AB]$', q_id_cell)
                or (len(cells) > 1 and re.search(r'kota\s+tempat\s+tinggal', cells[1], re.I))
            )
            if is_kota_row and len(cells) > 1:
                for item in _parse_cell(cells[1]):
                    if item[1].upper() not in seen_labels:
                        seen_labels.add(item[1].upper())
                        found.append(item)

    # Pass 2: fallback — scan seluruh sel yang mengandung ≥3 nama kota
    if len(found) < 3:
        found.clear(); seen_labels.clear()
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    txt = cell.text.strip()
                    if len(kota_re.findall(txt)) >= 3:
                        for item in _parse_cell(txt):
                            if item[1].upper() not in seen_labels:
                                seen_labels.add(item[1].upper())
                                found.append(item)
                        if len(found) >= 3:
                            break
                if len(found) >= 3: break
            if len(found) >= 3: break

    if found:
        try:
            found.sort(key=lambda x: int(x[0]))
        except ValueError:
            pass
        return found

    # Fallback: 11 kota standar Deka Jawa
    return [
        ('1', 'Jakarta'), ('2', 'Bogor'), ('3', 'Depok'), ('4', 'Tangerang'),
        ('5', 'Bekasi'), ('6', 'Bandung'), ('7', 'Cirebon'), ('8', 'Semarang'),
        ('9', 'Tegal'), ('10', 'Surabaya'), ('11', 'Madiun'),
    ]

def _extract_respondent_info(doc: docx.Document) -> list[dict]:
    """
    Ekstrak blok Informasi Responden dari halaman pertama dokumen.
    Menghasilkan list questions yang akan dijadikan begin_group info_responden.
    """
    info_tbl = None
    for tbl in doc.tables:
        if _is_info_table(tbl):
            info_tbl = tbl
            break

    # Field-field standar yang selalu ada di kuesioner Deka Research
    # Field ini akan dihasilkan bahkan jika tabel tidak ditemukan
    standard_fields = [
        ("nama_responden",    "Nama Responden",             "yes"),
        ("nomor_kuesioner",   "Nomor Kuesioner",            "yes"),
        ("alamat",            "Alamat Lengkap",             "no"),
        ("kelurahan",         "Kelurahan",                  "no"),
        ("kecamatan",         "Kecamatan",                  "no"),
        ("rt",                "RT",                         "no"),
        ("rw",                "RW",                         "no"),
        ("kota",              "Kota",                       "yes"),
        ("telp_rumah",        "Nomor Telepon Rumah",        "no"),
        ("telp_hp",           "Nomor HP / Telepon Kantor",  "no"),
        ("email",             "Alamat Email",               "no"),
        ("nama_interviewer",  "Nama Interviewer",           "yes"),
        ("no_interviewer",    "No. Interviewer",            "yes"),
        ("tanggal_wawancara", "Tanggal Wawancara",          "yes"),
        ("waktu_mulai",       "Waktu Mulai",                "yes"),
        ("waktu_selesai",     "Waktu Selesai",              "no"),
    ]

    # Jika tabel info ditemukan, coba ekstrak field tambahan (misal: nama toko, alamat toko)
    extra_fields: list[tuple] = []
    if info_tbl is not None:
        known_labels = {f[1].upper() for f in standard_fields}
        for row in info_tbl.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                # Cari label yang ada tanda : (format "Label :")
                if ":" in txt:
                    label_raw = txt.split(":")[0].strip()
                    if (label_raw and len(label_raw) > 3
                            and label_raw.upper() not in known_labels
                            and not any(m in label_raw.upper() for m in
                                        ["TANDA TANGAN", "DIPERIKSA", "SPV", "TL", "ESOMAR"])):
                        # Buat field_name dari label
                        fname = re.sub(r'[^a-z0-9]', '_',
                                       label_raw.lower()).strip('_')
                        fname = re.sub(r'_+', '_', fname)
                        if fname and len(fname) > 2:
                            lbl_upper = label_raw.upper()
                            if lbl_upper not in known_labels:
                                known_labels.add(lbl_upper)
                                extra_fields.append((fname, label_raw, "no"))

    # Tentukan tipe per field
    date_fields  = {"tanggal_wawancara"}
    time_fields  = {"waktu_mulai", "waktu_selesai"}
    kota_fields  = {"kota"}
    int_fields   = {"rt", "rw"}

    def _field_type(fname: str) -> str:
        if fname in date_fields:  return "date"
        if fname in time_fields:  return "time"
        if fname in kota_fields:  return "select_one list_kota"
        if fname in int_fields:   return "integer"
        return "text"

    questions: list[dict] = []

    # begin_group
    questions.append({
        "id": "info_responden",
        "is_header": False,
        "section": "",
        "subsection": "",
        "question": "INFORMASI RESPONDEN",
        "hint": "",
        "route_type": "begin_group",
        "choices": [],
        "skip_logic": "",
    })

    all_fields = standard_fields + extra_fields
    for fname, flabel, freq in all_fields:
        questions.append({
            "id": fname,
            "is_header": False,
            "section": "INFORMASI RESPONDEN",
            "subsection": "",
            "question": flabel,
            "hint": "",
            "route_type": _field_type(fname),
            "choices": [],
            "skip_logic": "",
            "_required": freq,
        })

    # end_group
    questions.append({
        "id": "end_info_responden",
        "is_header": False,
        "section": "",
        "subsection": "",
        "question": "",
        "hint": "",
        "route_type": "end_group",
        "choices": [],
        "skip_logic": "",
    })

    return questions


def _find_questionnaire_tables(doc: docx.Document) -> tuple[Table | None, Table | None]:
    screening_tbl = main_tbl = None
    for tbl in doc.tables:
        if screening_tbl is None and _is_screening_table(tbl):
            screening_tbl = tbl
        elif main_tbl is None and _is_main_table(tbl):
            main_tbl = tbl
        if screening_tbl and main_tbl:
            break
    return screening_tbl, main_tbl


# ── Special parsers ────────────────────────────────────────────────────────────

def _parse_awareness_block(q_id_prefix: str, cell: _Cell,
                           section: str, subsection: str,
                           merek_list: list[tuple],
                           product_name: str) -> list[dict]:
    """
    Parse blok awareness multi-sub (A1, A2, Y1).

    v10 fix: ekstrak teks sub-pertanyaan dari raw teks dokumen, BUKAN hardcode label.
    Sub-pertanyaan a/b/c/d/e/f ada inline dalam teks satu sel, formatnya:
      "a. Apakah merek ... (SA)\nb. Merek ... apalagi? (MA)\nc. ... (MA)\n..."
    atau kadang berlabel "A1a", "A1b", dst.

    Jika tidak ditemukan pattern inline, fallback ke label generik.
    """
    raw = _cell_text(cell)
    q_text, hint = _split_hint(raw)

    merek_choices = _choices_from_list(merek_list)

    # ── Coba ekstrak sub-pertanyaan inline dari teks ─────────────────────────
    # Pattern: "a." atau "a)" atau "A1a" diikuti teks pertanyaan
    # Juga handle format "a. Teks (SA)" atau "a. Teks\nProbe"
    inline_subs: list[tuple[str, str, str]] = []  # (suffix, teks, route_hint)

    # Pattern 1: "a. Teks..." dengan baris terpisah
    for m in re.finditer(
        r'(?:^|\n)\s*([a-f])[.)\s]+([^\n]{10,200})',
        raw, re.MULTILINE | re.IGNORECASE
    ):
        suffix = m.group(1).lower()
        teks   = m.group(2).strip()
        # Pastikan suffix berurutan (a, b, c, ...)
        expected = chr(ord('a') + len(inline_subs))
        if suffix == expected:
            inline_subs.append((suffix, teks, ""))

    # Pattern 2: "A1a Teks..." atau "A2a Teks..." dengan prefix q_id
    if not inline_subs:
        for m in re.finditer(
            r'(?:^|\n)\s*' + re.escape(q_id_prefix.upper()) + r'([a-f1])\s+([^\n]{10,200})',
            raw, re.MULTILINE | re.IGNORECASE
        ):
            suffix = m.group(1).lower()
            teks   = m.group(2).strip()
            inline_subs.append((suffix, teks, ""))

    # ── Deteksi route per sub dari tanda (SA)/(MA) di teks ───────────────────
    def _detect_route(teks: str) -> str:
        if re.search(r'\(SA\)', teks, re.I): return "SA"
        if re.search(r'\(MA\)', teks, re.I): return "MA"
        if re.search(r'\(M\)', teks, re.I):  return "MA"
        return ""  # akan di-assign default di bawah

    # ── Ambil choices sumber tahu dari nested table ───────────────────────────
    sumber_choices: list[dict] = []
    for nt in cell.tables:
        if len(nt.columns) >= 4:
            for row in nt.rows[1:]:
                cells_nt = row.cells
                label = cells_nt[0].text.strip()
                code  = cells_nt[1].text.strip() if len(cells_nt) > 1 else ""
                if label and not re.match(r'^[A-Z]\d+[a-z]?$', label) and code:
                    sumber_choices.append({"code": code, "label": label, "routing": ""})
            if sumber_choices:
                break

    # ── Cek apakah ada sub "f" (sumber tahu) ─────────────────────────────────
    has_f = bool(sumber_choices) or bool(
        re.search(r'sumber\s+tahu|dari\s+mana.{0,50}mengetahui|\bf\b.{0,10}(SA|MA)', raw, re.I)
    )
    # Cek dari header nested table
    for nt in cell.tables:
        h = [c.text.strip().upper() for c in nt.rows[0].cells] if nt.rows else []
        if any(q_id_prefix.upper() + "F" in hh or "SUMBER" in hh for hh in h):
            has_f = True
            break

    # ── Fallback: definisi default jika tidak ada inline ─────────────────────
    default_sub_defs = [
        ("a", "SA",  f"Merek {product_name} yang pertama kali Anda ingat? (TOM)",
                     "SPONTAN – satu jawaban"),
        ("b", "MA",  f"Merek {product_name} lain yang Anda ketahui?",
                     "SPONTAN – semua yang disebutkan"),
        ("c", "MA",  f"Merek {product_name} mana yang Anda ketahui? (Dibantu)",
                     "TUNJUKKAN KARTU BANTU"),
        ("d", "SA",  f"Merek {product_name} yang paling Anda favoritkan?",
                     "Satu merek"),
        ("e", "SA",  f"Merek {product_name} favorit kedua Anda?",
                     "Satu merek"),
    ]
    if has_f:
        default_sub_defs.append(("f", "MA",
            f"Dari mana Anda mengetahui merek {product_name} tersebut? (Sumber tahu)",
            "BACAKAN – boleh pilih lebih dari satu"))

    # ── Bangun list sub-pertanyaan final ──────────────────────────────────────
    # Jika berhasil ekstrak inline, gunakan teks asli dari dokumen
    # dengan route dari tanda (SA)/(MA) di teks, fallback ke default route
    default_routes = {"a":"SA","b":"MA","c":"MA","d":"SA","e":"SA","f":"MA"}

    if inline_subs and len(inline_subs) >= 3:
        # Pakai teks dari dokumen
        sub_defs_final = []
        for suffix, teks, _ in inline_subs:
            route = _detect_route(teks) or default_routes.get(suffix, "SA")
            # Bersihkan tanda (SA)/(MA) dari label pertanyaan
            teks_clean = re.sub(r'\s*\((?:SA|MA|M|OE)\)\s*', ' ', teks).strip()
            sub_defs_final.append((suffix, route, teks_clean, ""))
        # Tambah sub "f" jika ada sumber tahu tapi belum ter-ekstrak inline
        if has_f and not any(s[0] == "f" for s in sub_defs_final):
            sub_defs_final.append(("f", "MA",
                f"Dari mana Anda mengetahui merek {product_name} tersebut?", ""))
    else:
        # Gunakan fallback default
        sub_defs_final = [(s,r,l,h) for s,r,l,h in default_sub_defs]

    # ── skip_logic ────────────────────────────────────────────────────────────
    skip_matches = re.findall(
        r'(?:TANYAKAN|DITANYAKAN)\s+(?:JIKA|KEPADA)\s+[^\.]{5,80}', raw, re.I
    )
    skip_logic = "; ".join(skip_matches)

    items = []
    for suffix, route, label, hint_sub in sub_defs_final:
        full_id = f"{q_id_prefix}{suffix}"
        ch = merek_choices if suffix in ('a','b','c','d','e') else sumber_choices
        items.append(_make_q(
            q_id=full_id, section=section, subsection=subsection,
            question=label, hint=hint_sub or hint, route_type=route,
            choices=ch, skip_logic=skip_logic if suffix == 'a' else "",
            raw_text=raw,
        ))

    return items


def _parse_usage_block(q_id_prefix: str, cell: _Cell,
                       section: str, subsection: str,
                       merek_list: list[tuple],
                       product_name: str) -> list[dict]:
    """
    Parse blok usage V/Z (V1–V5 atau Z1–Z5) + sub-pertanyaan BUMO (V4a–V4e / Z4a–Z4e).
    """
    raw = _cell_text(cell)
    skip_logic = ""
    items = []

    merek_choices = _choices_from_list(merek_list)
    usia_choices  = _choices_from_list(_USIA_RANGE)
    time_choices  = _choices_from_list(_TIME_RANGE)

    # Ambil choices alasan dari nested tables
    alasan_choices: list[dict] = []
    alasan_negatif_choices: list[dict] = []
    for nt in cell.tables:
        ch = _parse_simple_nested_table(nt)
        if not ch:
            continue
        first_label = ch[0].get("label","").lower()
        if "mudah didapat" in first_label or "terjangkau" in first_label:
            alasan_choices = ch
        elif "sulit didapat" in first_label or "mahal" in first_label:
            alasan_negatif_choices = ch

    p = q_id_prefix  # 'v' atau 'z'
    P = p.upper()

    sub_defs = [
        (f"{p}1",   "MA", f"Merek {product_name} apa saja yang pernah Anda gunakan?",
                    "TUNJUKKAN KARTU BANTU", merek_choices),
        (f"{p}1a",  "OE", f"Mengapa Anda menggunakan merek-merek {product_name} tersebut?",
                    "PROBE", alasan_choices),
        (f"{p}2",   "MA", f"Merek {product_name} apa yang Anda gunakan dalam 1 tahun terakhir?",
                    "TUNJUKKAN KARTU BANTU", merek_choices),
        (f"{p}2a",  "OE", f"Mengapa Anda menggunakan merek-merek {product_name} tersebut (1 tahun terakhir)?",
                    "PROBE", alasan_choices),
        (f"{p}3",   "MA", f"Merek {product_name} apa yang Anda gunakan terakhir kali?",
                    "TUNJUKKAN KARTU BANTU", merek_choices),
        (f"{p}3a",  "OE", f"Mengapa Anda menggunakan merek {product_name} tersebut (terakhir kali)?",
                    "PROBE", alasan_choices),
        (f"{p}4a",  "SA", f"Merek {product_name} yang paling sering Anda gunakan saat ini? (BUMO)",
                    "TUNJUKKAN KARTU BANTU – satu merek", merek_choices),
        (f"{p}4b",  "SA", f"Sejak usia berapa pertama kali Anda membeli {product_name} merek tersebut?",
                    "PILIH RANGE USIA", usia_choices),
        (f"{p}4c",  "SA", f"Merek {product_name} apa yang Anda gunakan sebelum merek saat ini?",
                    "TUNJUKKAN KARTU BANTU", merek_choices),
        (f"{p}4c1", "MA", f"Mengapa Anda beralih dari merek sebelumnya ke merek saat ini ({product_name})?",
                    "PROBE – semua alasan", alasan_choices),
        (f"{p}4d",  "SA", f"Kapan pertama kali Anda beralih ke merek {product_name} saat ini?",
                    "PILIH RANGE WAKTU", time_choices),
        (f"{p}4e",  "OE", f"Apa yang mendorong Anda pertama kali mencoba merek {product_name} saat ini?",
                    "SPONTAN", []),
        (f"{p}5",   "SA", f"Merek {product_name} terbaik menurut Anda?",
                    "TUNJUKKAN KARTU BANTU – satu merek", merek_choices),
    ]

    for sub_id, route, label, hint_sub, choices in sub_defs:
        items.append(_make_q(
            q_id=sub_id, section=section, subsection=subsection,
            question=label, hint=hint_sub, route_type=route,
            choices=choices, raw_text=raw,
        ))

    return items


def _parse_q4_block(cell: _Cell, section: str, subsection: str,
                    merek_list: list[tuple]) -> list[dict]:
    """Pecah Q4 (BUMO produk utama) menjadi Q4a–Q4e."""
    raw = _cell_text(cell)
    merek_choices = _choices_from_list(merek_list)
    usia_choices  = _choices_from_list(_USIA_RANGE)
    time_choices  = _choices_from_list(_TIME_RANGE)

    # Ambil choices alasan ganti dari nested table (yang terpanjang)
    alasan_choices: list[dict] = []
    for nt in cell.tables:
        ch = _parse_simple_nested_table(nt)
        if len(ch) > len(alasan_choices):
            alasan_choices = ch

    items = [
        _make_q("q4a",  section, subsection,
                f"Merek {_PRODUCT_LABELS["primary"]} utama yang paling sering Anda gunakan saat ini? (BUMO)",
                "TUNJUKKAN KARTU BANTU – satu merek", "SA", merek_choices, raw_text=raw),
        _make_q("q4b",  section, subsection,
                f"Sejak usia berapa pertama kali Anda membeli {_PRODUCT_LABELS["primary"]} merek tersebut?",
                "PILIH RANGE USIA", "SA", usia_choices, raw_text=raw),
        _make_q("q4c",  section, subsection,
                f"Merek {_PRODUCT_LABELS["primary"]} apa yang Anda gunakan sebelum merek saat ini?",
                "TUNJUKKAN KARTU BANTU", "SA", merek_choices, raw_text=raw),
        _make_q("q4c1", section, subsection,
                f"Mengapa Anda beralih dari merek {_PRODUCT_LABELS["primary"]} sebelumnya ke merek saat ini?",
                "PROBE – semua alasan", "MA", alasan_choices, raw_text=raw),
        _make_q("q4d",  section, subsection,
                f"Kapan pertama kali Anda beralih ke merek {_PRODUCT_LABELS["primary"]} saat ini?",
                "PILIH RANGE WAKTU", "SA", time_choices, raw_text=raw),
        _make_q("q4e",  section, subsection,
                f"Apa yang mendorong Anda pertama kali mencoba merek {_PRODUCT_LABELS["primary"]} saat ini?",
                "SPONTAN – terbuka", "OE", [], raw_text=raw),
    ]
    return items


def _parse_md1_block(cell: _Cell, section: str, subsection: str) -> list[dict]:
    """Pecah MD1 (4 sub-kolom media sosial) → MD1a, MD1b, MD1c, MD1d."""
    raw = _cell_text(cell)
    # Ambil choices dari nested table 5-kolom (header + 4 sub)
    platform_choices: list[dict] = []
    for nt in cell.tables:
        if len(nt.columns) >= 4:
            for row in nt.rows[2:]:  # skip 2 header rows
                cells = row.cells
                label = cells[0].text.strip()
                code  = cells[1].text.strip() if len(cells) > 1 else ""
                if label and not re.match(r'^MD\d', label) and code:
                    platform_choices.append({"code": code, "label": label, "routing": ""})
            break

    sub_defs = [
        ("md1a", "Sosial media yang Anda ketahui?",               "SPONTAN + TUNJUKKAN KARTU"),
        ("md1b", "Sosial media yang Anda akses dalam 3 bulan terakhir?", "P3M"),
        ("md1c", "Sosial media yang Anda akses dalam 1 minggu terakhir?", "P1W"),
        ("md1d", "Sosial media yang Anda akses setiap hari?",     "SETIAP HARI"),
    ]

    return [
        _make_q(sid, section, subsection, label, hint, "MA",
                platform_choices, raw_text=raw)
        for sid, label, hint in sub_defs
    ]


def _parse_matrix_q(q_id: str, cell: _Cell, section: str, subsection: str,
                    merek_list: list[tuple], scale_type: str = "10") -> list[dict]:
    """
    Parse matrix pertanyaan (Q9b, Q10, Q14c, Q15a) → begin_group + per-atribut/merek questions.
    Untuk Q14c/Q15a: per-atribut select_one list_merek
    Untuk Q9b/Q10: per-atribut integer (skala 1-10)
    """
    raw = _cell_text(cell)
    q_text, hint = _split_hint(raw)
    items = []

    atribut_list: list[str] = []
    merek_choices = _choices_from_list(merek_list) if merek_list else []

    # Ambil daftar atribut dari nested table
    for nt in cell.tables:
        rows = list(nt.rows)
        if not rows:
            continue
        nc = len(nt.columns)
        if nc >= 4:  # matrix banyak kolom = tabel atribut vs merek
            for row in rows[1:]:  # skip header
                atribut = row.cells[0].text.strip()
                if atribut and not re.match(r'^\[|^ROTASIKAN|^DP:', atribut, re.I):
                    atribut_list.append(atribut)
        elif nc == 2:  # tabel atribut dua kolom
            for row in rows[1:]:
                atribut = row.cells[0].text.strip()
                if atribut and not re.match(r'^\[|^ROTASIKAN', atribut, re.I):
                    atribut_list.append(atribut)

    if not atribut_list:
        # Fallback: buat satu field saja
        return [_make_q(q_id, section, subsection, q_text, hint,
                        "", [], raw_text=raw)]

    # begin_group
    grp_id = f"grp_{q_id.lower()}"
    items.append(_make_q(grp_id, section, subsection,
                         q_text, hint, "begin_group", [], raw_text=raw))

    skala_choices = _choices_from_list(_SKALA_10 if scale_type == "10" else _SKALA_11)

    for i, atribut in enumerate(atribut_list, 1):
        safe_attr = re.sub(r'[^a-z0-9]', '_', atribut.lower())[:25].strip('_')
        sub_id = f"{q_id.lower()}_{i:02d}"
        if merek_list and q_id.upper() in ("Q14C", "Q15A"):
            # Matrix atribut vs merek → select_one list_merek per atribut
            items.append(_make_q(
                sub_id, section, subsection,
                question=f"{atribut} — merek mana yang paling sesuai?",
                hint=atribut, route_type="SA",
                choices=merek_choices, raw_text=atribut,
            ))
        else:
            # Rating skala 1-10 per atribut
            items.append(_make_q(
                sub_id, section, subsection,
                question=atribut, hint=hint,
                route_type="SA", choices=skala_choices, raw_text=atribut,
            ))

    items.append(_make_q(f"{grp_id}_end", section, subsection,
                         "", "", "end_group", [], raw_text=""))
    return items


def _parse_q14a_block(q_id: str, cell: _Cell,
                      section: str, subsection: str) -> list[dict]:
    """Q14a: per-warna pilih merek. Q14b: per-bentuk logo pilih merek."""
    raw = _cell_text(cell)
    q_text, hint = _split_hint(raw)
    merek_choices = _choices_from_list(_get_merek("primary"))
    items = []

    atribut_list = []
    for nt in cell.tables:
        for row in nt.rows[1:]:  # skip header
            val = row.cells[0].text.strip()
            if val and not re.match(r'^ROTASIKAN', val, re.I):
                atribut_list.append(val)

    if not atribut_list:
        return [_make_q(q_id, section, subsection, q_text, hint, "SA",
                        merek_choices, raw_text=raw)]

    items.append(_make_q(f"grp_{q_id.lower()}", section, subsection,
                         q_text, hint, "begin_group", [], raw_text=raw))
    for i, attr in enumerate(atribut_list, 1):
        items.append(_make_q(
            f"{q_id.lower()}_{i:02d}", section, subsection,
            f"{q_text} — {attr}", attr, "SA", merek_choices, raw_text=attr,
        ))
    items.append(_make_q(f"grp_{q_id.lower()}_end", section, subsection,
                         "", "", "end_group", [], raw_text=""))
    return items


def _parse_q15b_block(cell: _Cell, section: str, subsection: str) -> list[dict]:
    """Q15b: ranking 10 merek PVC dari tertua ke termuda → ranking integer per merek."""
    raw = _cell_text(cell)
    q_text, hint = _split_hint(raw)
    merek_choices = _choices_from_list(_get_merek("primary")[:10])  # tanpa TT/TA

    items = [
        _make_q("grp_q15b", section, subsection,
                f"Urutkan merek {_PRODUCT_LABELS["primary"]} dari yang paling tua sampai paling baru",
                hint, "begin_group", [], raw_text=raw)
    ]
    for code, label in _get_merek("primary")[:10]:
        items.append(_make_q(
            f"q15b_{code}", section, subsection,
            f"Peringkat ke berapa merek {label}? (1=tertua, 10=termuda)",
            "Masukkan angka 1–10", "integer", [], raw_text=label,
        ))
    items.append(_make_q("grp_q15b_end", section, subsection,
                         "", "", "end_group", [], raw_text=""))
    return items


def _parse_slogan_q(q_id: str, cell: _Cell, section: str, subsection: str,
                    merek_list: list[tuple]) -> list[dict]:
    """Qxxb 'Untuk merek apa slogan tersebut?' → select_one list_merek."""
    raw = _cell_text(cell)
    q_text, hint = _split_hint(raw)
    if not q_text or "untuk merek apa" not in q_text.lower():
        q_text = "Untuk merek apa slogan tersebut?"
    merek_choices = _choices_from_list(merek_list)
    return [_make_q(q_id, section, subsection, q_text, hint, "SA",
                    merek_choices, raw_text=raw)]


def _parse_m1_block(cell: _Cell, section: str, subsection: str) -> list[dict]:
    """
    M1: matrix media habit.
    M1a: pernah melakukan (MA) per kategori media
    M1b: frekuensi (SA skala frekuensi) per kategori media
    """
    raw = _cell_text(cell)
    q_text, hint = _split_hint(raw)

    media_list: list[tuple] = []
    freq_choices: list[dict] = []

    for nt in cell.tables:
        if len(nt.columns) >= 8:
            rows = list(nt.rows)
            # Header baris 2 = frekuensi labels
            if len(rows) > 1:
                freq_labels = [c.text.strip() for c in rows[1].cells[3:]]
                freq_choices = [
                    {"code": str(i+1), "label": lbl, "routing": ""}
                    for i, lbl in enumerate(freq_labels) if lbl
                ]
            # Data baris
            for row in rows[2:]:
                cells = row.cells
                media_label = cells[0].text.strip()
                code = cells[2].text.strip()  # M1a code
                if media_label and not re.match(r'^ROTASIKAN', media_label, re.I):
                    media_list.append((code or str(len(media_list)+1), media_label))
            break

    if not media_list:
        return [_make_q("m1", section, subsection, q_text, hint,
                        "MA", [], raw_text=raw)]

    items = [
        _make_q("grp_m1", section, subsection, q_text, hint,
                "begin_group", [], raw_text=raw)
    ]

    # M1a: select_multiple – kegiatan yang pernah dilakukan
    m1a_choices = _choices_from_list(media_list)
    items.append(_make_q("m1a", section, subsection,
                         "Kegiatan media apa saja yang pernah Anda lakukan?",
                         "MA – pilih semua yang berlaku", "MA", m1a_choices, raw_text=raw))

    # M1b: per-media yang dipilih di M1a → frekuensi
    for code, media_label in media_list:
        safe = re.sub(r'[^a-z0-9]', '_', media_label.lower())[:20].strip('_')
        items.append(_make_q(
            f"m1b_{safe}", section, subsection,
            f"Seberapa sering Anda {media_label.lower()}?",
            "PILIH SATU FREKUENSI", "SA",
            freq_choices,
            skip_logic=f"TANYAKAN JIKA m1a TERKODE {code}",
            raw_text=media_label,
        ))

    items.append(_make_q("grp_m1_end", section, subsection,
                         "", "", "end_group", [], raw_text=""))
    return items


def _parse_rating_q(q_id: str, cell: _Cell, section: str, subsection: str,
                    scale_size: int = 10) -> list[dict]:
    """Rating pertanyaan (skala 1-10 atau NPS 0-10) → select_one."""
    raw = _cell_text(cell)
    q_text, hint = _split_hint(raw)
    skip_matches = re.findall(
        r'(?:TANYAKAN|DITANYAKAN)\s+(?:JIKA|KEPADA)\s+[^\.]{5,80}', raw, re.I
    )
    skip_logic = "; ".join(skip_matches)
    choices = _choices_from_list(_SKALA_11 if scale_size == 11 else _SKALA_10)
    return [_make_q(q_id, section, subsection, q_text, hint,
                    "SA", choices, skip_logic=skip_logic, raw_text=raw)]


def _parse_brand_statement_matrix(q_id: str, cell: _Cell,
                                   section: str, subsection: str) -> list[dict]:
    """
    Parse Q8a/Q8b/Q8c/Q8d — matrix merek vs pernyataan.
    Struktur NT: baris 0 = ['ROTASIKAN MEREK', 'Alderon', 'Intilon', ...]
                 baris 1+ = ['Pernyataan...', '1', '1', ...]
    Dikonversi ke: begin_group + per-pernyataan select_multiple list_merek
    """
    raw = _cell_text(cell)
    q_text, hint = _split_hint(raw)
    items = []

    merek_list: list[tuple] = []
    stmt_list: list[str] = []

    for nt in cell.tables:
        rows = list(nt.rows)
        if not rows or len(nt.columns) < 8:
            continue
        # Baris 0: header merek
        h = [c.text.strip() for c in rows[0].cells]
        if any(re.match(r'^(ROTASIKAN|Alderon|Intilon|Grest)', hh, re.I) for hh in h):
            for idx, label in enumerate(h[1:], 1):
                if label and not re.match(r'^ROTASIKAN', label, re.I):
                    merek_list.append((str(idx), label))
            # Baris 1+: pernyataan
            for row in rows[1:]:
                stmt = row.cells[0].text.strip()
                if stmt and not re.match(r'^ROTASIKAN|^DP:', stmt, re.I):
                    stmt_list.append(stmt)
            break

    if not merek_list or not stmt_list:
        # Fallback: satu pertanyaan SA list merek
        return [_make_q(q_id, section, subsection, q_text, hint,
                        "SA", _choices_from_list(_get_merek("primary")), raw_text=raw)]

    merek_choices = _choices_from_list(merek_list)
    items.append(_make_q(f"grp_{q_id}", section, subsection,
                         q_text, hint, "begin_group", [], raw_text=raw))
    for i, stmt in enumerate(stmt_list, 1):
        safe = re.sub(r'[^a-z0-9]', '_', stmt.lower())[:20].strip('_')
        items.append(_make_q(
            f"{q_id}_{i:02d}", section, subsection,
            stmt, hint, "SA", merek_choices, raw_text=stmt,
        ))
    items.append(_make_q(f"grp_{q_id}_end", section, subsection,
                         "", "", "end_group", [], raw_text=""))
    return items


def _parse_media_matrix_q(q_id: str, cell: _Cell,
                           section: str, subsection: str) -> list[dict]:
    """
    Parse M12/MD8 — matrix per-media dengan skala 1-10.
    Struktur NT: baris 0 = header (label skala), baris 1+ = ['Media label', kategori, 1..10]
    Dikonversi ke: begin_group + per-media integer/select_one skala
    """
    raw = _cell_text(cell)
    q_text, hint = _split_hint(raw)
    skala_choices = _choices_from_list(_SKALA_10)
    items = []

    media_list: list[tuple] = []  # (safe_id, label)

    for nt in cell.tables:
        rows = list(nt.rows)
        if not rows or len(nt.columns) < 10:
            continue
        for row in rows[1:]:  # skip header
            cells_nt = [c.text.strip() for c in row.cells]
            label = cells_nt[0]
            if label and not re.match(r'^ROTASIKAN|^DP:|^$', label, re.I):
                safe = re.sub(r'[^a-z0-9]', '_', label.lower())[:20].strip('_')
                media_list.append((safe, label))
        break

    if not media_list:
        return [_make_q(q_id, section, subsection, q_text, hint,
                        "SA", skala_choices, raw_text=raw)]

    items.append(_make_q(f"grp_{q_id}", section, subsection,
                         q_text, hint, "begin_group", [], raw_text=raw))
    for safe, label in media_list:
        items.append(_make_q(
            f"{q_id}_{safe}", section, subsection,
            f"{label} — {q_text.lower()}", hint,
            "SA", skala_choices, raw_text=label,
        ))
    items.append(_make_q(f"grp_{q_id}_end", section, subsection,
                         "", "", "end_group", [], raw_text=""))
    return items
    """Rating pertanyaan (V7, V8, V9, V10, Z7, Z8, V11, dll) → integer atau select_one."""
    raw = _cell_text(cell)
    q_text, hint = _split_hint(raw)
    skip_matches = re.findall(
        r'(?:TANYAKAN|DITANYAKAN)\s+(?:JIKA|KEPADA)\s+[^\.]{5,80}', raw, re.I
    )
    skip_logic = "; ".join(skip_matches)

    # NPS 0-10 → select_one dengan 11 pilihan
    if scale_size == 11:
        choices = _choices_from_list(_SKALA_11)
        return [_make_q(q_id, section, subsection, q_text, hint,
                        "SA", choices, skip_logic=skip_logic, raw_text=raw)]
    else:
        choices = _choices_from_list(_SKALA_10)
        return [_make_q(q_id, section, subsection, q_text, hint,
                        "SA", choices, skip_logic=skip_logic, raw_text=raw)]



def _extract_inline_subs(q_id: str, cell: _Cell,
                         section: str, subsection: str) -> list[dict] | None:
    """
    v10: Deteksi sub-pertanyaan inline (a./b./c./d. dalam satu sel).

    Jika sel berisi pola "a. Teks... (SA/MA/OE)\nb. Teks... (SA/MA)\n..."
    maka hasilkan baris terpisah Q_IDa, Q_IDb, Q_IDc, dst.

    Returns None jika tidak ada pola inline yang ditemukan.
    Returns list[dict] jika ditemukan ≥2 sub.
    """
    raw = _cell_text(cell)

    # Cari semua sub-pertanyaan dengan pola "a. Teks" atau "a) Teks"
    # Minimal 2 sub agar dianggap pattern inline
    subs: list[tuple[str, str]] = []  # (suffix, full_teks_sub)

    # Match pola: (a/b/c/..) diikuti titik/kurung dan teks, diakhiri newline atau sub berikutnya
    pattern = re.compile(
        r'(?:^|\n)[ \t]*([a-f])[.)][\s]+(.+?)(?=(?:\n[ \t]*[a-f][.)][\s])|$)',
        re.DOTALL | re.MULTILINE
    )
    for m in pattern.finditer(raw):
        suffix = m.group(1).lower()
        teks   = m.group(2).strip().replace("\n", " ")
        # Harus berurutan
        expected = chr(ord('a') + len(subs))
        if suffix != expected:
            break
        subs.append((suffix, teks))

    if len(subs) < 2:
        return None

    # Deteksi route dari tanda (SA)/(MA)/(OE) di teks
    def _route(teks: str) -> str:
        if re.search(r'\(SA\)', teks, re.I): return "SA"
        if re.search(r'\(MA\)', teks, re.I): return "MA"
        if re.search(r'\b(MA|M)\b', teks): return "MA"
        if re.search(r'\(OE\)', teks, re.I): return "OE"
        return "SA"  # default

    items = []
    for suffix, teks in subs:
        full_id = f"{q_id.lower()}{suffix}"
        route = _route(teks)
        # Bersihkan tanda route dari label
        label = re.sub(r'\s*\((?:SA|MA|M|OE)\)\s*', ' ', teks).strip()
        q_text_clean, hint_clean = _split_hint(label)
        items.append(_make_q(
            q_id=full_id, section=section, subsection=subsection,
            question=q_text_clean, hint=hint_clean, route_type=route,
            choices=[], raw_text=teks,
        ))

    return items

# ── Main row parser ────────────────────────────────────────────────────────────

def _parse_question_row(q_id: str, cell: _Cell, route_raw: str,
                        section: str, subsection: str) -> list[dict]:
    """
    Parse satu baris → bisa menghasilkan 1 atau banyak pertanyaan (sub-pertanyaan).
    """
    raw_text  = _cell_text(cell)
    q_text, hint_text = _split_hint(raw_text)
    route_type = _normalize_route(route_raw)

    q_id_upper = q_id.upper()

    # ── Awareness blocks ──────────────────────────────────────────────────────
    if q_id_upper == "A1":
        return _parse_awareness_block("a1", cell, section, subsection,
                                      _get_merek("primary"), _PRODUCT_LABELS["primary"])
    if q_id_upper == "A2":
        return _parse_awareness_block("a2", cell, section, subsection,
                                      _get_merek("secondary"), _PRODUCT_LABELS["secondary"])
    if q_id_upper == "Y1":
        return _parse_awareness_block("y1", cell, section, subsection,
                                      _get_merek("tertiary"), _PRODUCT_LABELS["tertiary"])

    # ── Usage blocks (V/Z) ───────────────────────────────────────────────────
    if q_id_upper == "V":
        return _parse_usage_block("v", cell, section, subsection,
                                  _get_merek("secondary"), _PRODUCT_LABELS["secondary"])
    if q_id_upper == "Z":
        return _parse_usage_block("z", cell, section, subsection,
                                  _get_merek("tertiary"), _PRODUCT_LABELS["tertiary"])

    # ── Q4 BUMO block ─────────────────────────────────────────────────────────
    if q_id_upper == "Q4":
        return _parse_q4_block(cell, section, subsection, _get_merek("primary"))

    # ── MD1 block ─────────────────────────────────────────────────────────────
    if q_id_upper == "MD1":
        return _parse_md1_block(cell, section, subsection)

    # ── Matrix pertanyaan ──────────────────────────────────────────────────────
    if q_id_upper in ("Q9B", "Q9b"):
        return _parse_matrix_q("Q9b", cell, section, subsection, [], "10")
    if q_id_upper == "Q10":
        return _parse_matrix_q("Q10", cell, section, subsection, [], "10")
    if q_id_upper in ("Q14C", "Q14c"):
        return _parse_matrix_q("Q14c", cell, section, subsection, _get_merek("primary"), "10")
    if q_id_upper in ("Q15A", "Q15a"):
        return _parse_matrix_q("Q15a", cell, section, subsection, _get_merek("primary"), "10")

    # ── Q14a, Q14b (warna/bentuk logo per merek) ──────────────────────────────
    if q_id_upper in ("Q14A", "Q14a"):
        return _parse_q14a_block("Q14a", cell, section, subsection)
    if q_id_upper in ("Q14B", "Q14b"):
        return _parse_q14a_block("Q14b", cell, section, subsection)

    # ── Q15b ranking merek ────────────────────────────────────────────────────
    if q_id_upper in ("Q15B", "Q15b"):
        return _parse_q15b_block(cell, section, subsection)

    # ── Slogan awareness Qxxa / Vxxa / Zxxa → SA ya/tidak ───────────────────
    # Pattern: Q16a, Q17a, Q22a, dst — "Apakah pernah mendengar/membaca slogan"
    slogan_aware_match = re.match(r'^([QVZ]\d{2,}[A-Z]*)A$', q_id_upper)
    if slogan_aware_match:
        q_text_lower = q_text.lower()
        if "slogan" in q_text_lower or "pernah mendengar" in q_text_lower or "pernah membaca" in q_text_lower:
            ya_tidak = [
                {"code": "1", "label": "Ya, pernah", "routing": ""},
                {"code": "2", "label": "Tidak pernah", "routing": ""},
            ]
            return [_make_q(q_id.lower(), section, subsection, q_text, hint_text,
                            "SA", ya_tidak, raw_text=raw_text)]

    # ── Slogan kesesuaian Qxxc / Vxxc / Zxxc → rating skala 1-10 ────────────
    slogan_rating_match = re.match(r'^([QVZ]\d{2,}[A-Z]*)C$', q_id_upper)
    if slogan_rating_match:
        q_text_lower = q_text.lower()
        if "sesuai" in q_text_lower or "skala" in q_text_lower or "slogan" in q_text_lower:
            return _parse_rating_q(q_id.lower(), cell, section, subsection, 10)

    # ── Slogan attribution Qxxb / Vxxb / Zxxb ────────────────────────────────
    # Fix: \d+ (dua atau lebih digit) agar Q16b, Q17b, dll tertangkap
    slogan_match = re.match(r'^([QVZ]\d+[A-Z]*)B$', q_id_upper)
    if slogan_match:
        q_text_lower = q_text.lower()
        if "slogan" in q_text_lower or "untuk merek apa" in q_text_lower:
            merek = _get_merek("secondary") if q_id_upper.startswith("V") else (
                    _get_merek("tertiary")     if q_id_upper.startswith("Z") else _get_merek("primary"))
            return _parse_slogan_q(q_id.lower(), cell, section, subsection, merek)

    # ── Q8a/Q8b/Q8c/Q8d — matrix merek (header=merek, baris=pernyataan) ──────
    # NT: 11 kolom, baris 0 = ['ROTASIKAN MEREK', 'Alderon', 'Intilon', ...]
    if q_id_upper in ("Q8A","Q8B","Q8C","Q8D"):
        return _parse_brand_statement_matrix(q_id.lower(), cell, section, subsection)

    # ── E5 — matrix atribut iklan (11-kolom skala per atribut) ───────────────
    if q_id_upper == "E5":
        return _parse_matrix_q("E5", cell, section, subsection, [], "10")

    # ── M12 — matrix media attention (12-kolom, per-media baris) ─────────────
    if q_id_upper == "M12":
        return _parse_media_matrix_q("m12", cell, section, subsection)

    # ── MD8 — matrix social media influence (11-kolom per-platform) ──────────
    if q_id_upper == "MD8":
        return _parse_media_matrix_q("md8", cell, section, subsection)

    # ── MD25 — OE (open-ended, 1x1 empty nested table) ───────────────────────
    if q_id_upper == "MD25":
        q_text2, hint2 = _split_hint(raw_text)
        return [_make_q(q_id.lower(), section, subsection, q_text2, hint2,
                        "OE", [], raw_text=raw_text)]

    # ── MD26a — matrix konten Youtube (11-kolom skala) ───────────────────────
    if q_id_upper == "MD26A":
        return _parse_matrix_q("MD26a", cell, section, subsection, [], "10")
        return _parse_m1_block(cell, section, subsection)

    # ── Rating 1-10 (V7, V8, V9, V10, V11, Z7, Z8, Z9, Z10, Z11c, V13c) ─────
    nps_ids  = {"V10","Z10","MD12","MD19","Q13","Q12A","Q12B"}
    rate_ids = {
        "V7","V8","V9","V11","Z7","Z8","Z9","Z11","V13C","Z11C",
        "Q11","Q16C","Q17C","Q18C","Q19C","Q20C","Q22C","Q23C",
        "E4","E6","E7","F3","F4","V13C","Z11C",
    }
    if q_id_upper in nps_ids:
        return _parse_rating_q(q_id.lower(), cell, section, subsection, 11)
    if q_id_upper in rate_ids:
        return _parse_rating_q(q_id.lower(), cell, section, subsection, 10)

    # ── Auto-detect skala dari nested table ───────────────────────────────────
    # Kalau ada nested table 10+ kolom dengan baris pertama berisi angka 1-10
    # → pasti skala rating
    for nt in cell.tables:
        nc = len(nt.columns)
        if nc < 10:
            continue
        rows_nt = list(nt.rows)
        # Cek baris terakhir (atau baris ke-2) berisi angka berurutan
        check_row = rows_nt[-1] if rows_nt else None
        if check_row:
            vals = [c.text.strip() for c in check_row.cells]
            nums = [v for v in vals if re.match(r'^\d+$', v)]
            if len(nums) >= 8:
                # Baris 0 mungkin label "Sangat tidak..." / "0"
                first_val = rows_nt[0].cells[0].text.strip()
                scale_size = 11 if "0" in [c.text.strip() for c in rows_nt[-1].cells] and nc >= 11 else 10
                return _parse_rating_q(q_id.lower(), cell, section, subsection, scale_size)

    # ── Sub-pertanyaan inline (a./b./c. dalam satu sel) ─────────────────────
    inline_result = _extract_inline_subs(q_id, cell, section, subsection)
    if inline_result and len(inline_result) >= 2:
        # Inject choices dari nested tables ke sub-pertanyaan jika ada
        all_ch: list[dict] = []
        for nt in cell.tables:
            if len(nt.columns) < 10:
                all_ch.extend(_parse_simple_nested_table(nt))
        all_ch = [c for c in all_ch
                  if not re.match(r'^\[DP|^ROTASIKAN|^DP:', c.get("label",""), re.I)]
        if all_ch:
            for q in inline_result:
                if not q.get("choices"):
                    q["choices"] = all_ch
        return inline_result

    # ── Standard question ─────────────────────────────────────────────────────
    all_choices: list[dict] = []
    for nt in cell.tables:
        # Skip tabel skala yang sudah tidak dihandle di atas
        if len(nt.columns) >= 10:
            continue
        all_choices.extend(_parse_simple_nested_table(nt))

    # Bersihkan choices yang merupakan instruksi DP
    all_choices = [
        c for c in all_choices
        if not re.match(r'^\[DP|^ROTASIKAN|^DP:', c.get("label",""), re.I)
        and not re.match(r'^\[DP|^ROTASIKAN', c.get("code",""), re.I)
    ]

    is_grid = bool(re.search(r'\bROTASIKAN\b', raw_text, re.I))

    skip_matches = re.findall(
        r'(?:TANYAKAN|DITANYAKAN)\s+(?:JIKA|KEPADA)\s+[^\.]{5,80}',
        raw_text, re.I
    )
    skip_logic = "; ".join(skip_matches).strip()

    return [_make_q(
        q_id=q_id, section=section, subsection=subsection,
        question=q_text, hint=hint_text,
        route_type=route_type, choices=all_choices,
        skip_logic=skip_logic, is_grid=is_grid, raw_text=raw_text,
    )]


# ── Table walker ───────────────────────────────────────────────────────────────

def _walk_table(tbl: Table, section_default: str) -> list[dict]:
    items: list[dict] = []
    section    = section_default
    subsection = ""

    # ── Pass 1: kumpulkan semua brand grid dari baris kosong ────────────────
    # Key: frozenset of q_ids → choices mapping
    # Struktur: dict { norm_q_id: [choices] }
    brand_grid_registry: dict[str, list[dict]] = {}

    # Kumpulkan juga S8 PHOTOCARD grid dari tabel screening
    # (S8a/S8b/S8c share satu nested table di row S8c)
    for row in tbl.rows:
        cells   = [c.text.strip() for c in row.cells]
        q_id    = cells[0]
        q_cell  = row.cells[1]

        # Row bukan pertanyaan (id kosong) tapi punya nested table
        if not q_id:
            for nt in q_cell.tables:
                if len(nt.columns) < 6 or len(nt.rows) < 3:
                    continue
                grid_data = _parse_brand_grid_table(nt)
                brand_grid_registry.update(grid_data)

        # Row pertanyaan yang punya nested table S8a/S8b/S8c (multi-kolom)
        if q_id:
            for nt in q_cell.tables:
                if len(nt.columns) >= 4 and len(nt.rows) >= 2:
                    # Cek apakah header berisi S8a/S8b/S8c atau xxx-pertanyaan lain
                    h0 = [c.text.strip().lower() for c in nt.rows[0].cells]
                    if any(re.match(r'^s8[abc]$', h) for h in h0):
                        grid_data = _parse_brand_grid_table(nt)
                        brand_grid_registry.update(grid_data)

    # ── Pass 2: parse pertanyaan dan inject choices dari registry ───────────
    for row in tbl.rows:
        cells  = [c.text.strip() for c in row.cells]
        q_id   = cells[0]
        route  = cells[2] if len(cells) > 2 else ""
        raw_q  = _cell_text(row.cells[1])

        if not q_id and not raw_q:
            continue

        if not q_id:
            text_up = raw_q.strip()
            if not text_up or _NON_SECTION_RE.match(text_up):
                continue
            if len(text_up) < 100:
                if text_up.upper() in ("SCREENING", "KUESIONER UTAMA"):
                    section    = text_up
                    subsection = ""
                else:
                    subsection = text_up
                items.append(_make_q(
                    "", section, subsection, text_up, "",
                    "info", [], is_header=True, raw_text=text_up,
                ))
            continue

        if not _Q_ID_RE.match(q_id):
            continue

        parsed = _parse_question_row(q_id, row.cells[1], route, section, subsection)

        # Inject choices dari brand grid registry jika belum ada
        for q in parsed:
            if not q.get("choices"):
                qid_norm = q["id"].lower().replace("-","_")
                # Coba juga tanpa normalisasi hyphen (Q6b-1 → q6b-1)
                qid_hyp  = q["id"].lower()
                chosen = (brand_grid_registry.get(qid_norm)
                          or brand_grid_registry.get(qid_hyp)
                          or brand_grid_registry.get(qid_norm.rstrip("_1").rstrip("_"))
                          or [])
                if chosen:
                    q["choices"] = chosen

            # Q6b-1 / V6b-1 / Z6b-1 → consideration set, pakai merek yang sama
            # dengan Q6a / V6a / Z6a (tidak masuk brand grid karena kolom berbeda)
            qid_up = q["id"].upper()
            if not q.get("choices"):
                if re.match(r'^Q6B-?1$', qid_up):
                    q["choices"] = _choices_from_list(_get_merek("primary"))
                elif re.match(r'^V6B-?1$', qid_up):
                    q["choices"] = _choices_from_list(_get_merek("secondary"))
                elif re.match(r'^Z6B-?1$', qid_up):
                    q["choices"] = _choices_from_list(_get_merek("tertiary"))
            if q["id"].lower() in ("s8a","s8b","s8c"):
                if not q.get("choices"):
                    q["choices"] = _choices_from_list(_CHOICES_PRODUK)

            # S8d, S8e → jenis pipa plastik
            if q["id"].lower() in ("s8d","s8e"):
                if not q.get("choices"):
                    q["choices"] = _choices_from_list(_CHOICES_JENIS_PRODUK)

            # E2, MD24 → 'MASUKAN LIST BRAND' → list merek produk utama
            if not q.get("choices") and re.search(
                r'MASUKAN LIST BRAND|LIST BRAND PIPA', q.get("raw_text",""), re.I
            ):
                q["choices"] = _choices_from_list(_get_merek("primary"))

            # P7b → merek pipa SA (disebutkan di instruksi DP)
            if q["id"].lower() == "p7b" and not q.get("choices"):
                q["choices"] = _choices_from_list(_get_merek("primary"))

            # E3a, E3b, E3c — 1x1 empty nested table = open-ended
            if q["id"].lower() in ("e3a","e3b","e3c"):
                q["choices"]    = []
                q["route_type"] = "OE"

            # E4, E6, E7, F3, F4, M12, MD8 — skala 1-10 tanpa choices
            # (sudah dihandle di _parse_rating_q via standard path, tapi jaga-jaga)
            if q["id"].lower() in ("e4","e6","e7","f3","f4") and not q.get("choices"):
                q["choices"]    = _choices_from_list(_SKALA_10)
                q["route_type"] = "SA"

        items.extend(parsed)

    return items


# ── Public API ─────────────────────────────────────────────────────────────────

def parse_questionnaire(docx_path: str) -> list[dict]:
    doc = docx.Document(docx_path)

    # ── Ekstrak brand list & PHOTOCARD choices langsung dari dokumen ──────────
    _extract_brand_lists_from_doc(doc)

    # ── Ekstrak daftar kota dari screening ────────────────────────────────────
    kota_choices = _extract_kota_from_doc(doc)

    # ── Ekstrak blok Informasi Responden (halaman pertama) ────────────────────
    info_questions = _extract_respondent_info(doc)

    # Sisipkan kota choices sebagai metadata di question pertama agar
    # json_to_xlsform bisa membacanya untuk mengisi list_kota
    # Dibungkus dalam question tipe "kota_choices_data"
    info_questions.insert(0, {
        "id": "_kota_choices_data",
        "is_header": True,
        "section": "",
        "subsection": "",
        "question": "",
        "hint": "",
        "route_type": "kota_choices_data",
        "choices": [{"code": c, "label": l, "routing": ""} for c, l in kota_choices],
        "skip_logic": "",
    })

    # ── Temukan tabel screening & utama ───────────────────────────────────────
    screening_tbl, main_tbl = _find_questionnaire_tables(doc)

    if screening_tbl is None and main_tbl is None:
        raise ValueError(
            "Tidak ditemukan tabel kuesioner 3-kolom. "
            "Pastikan file adalah kuesioner Deka Research format tabel Word."
        )

    questions: list[dict] = []

    # Blok 1: Informasi Responden (selalu pertama)
    questions.extend(info_questions)

    # Blok 2: Screening
    if screening_tbl:
        questions.extend(_walk_table(screening_tbl, "SCREENING"))

    # Blok 3: Kuesioner Utama
    if main_tbl:
        questions.extend(_walk_table(main_tbl, "KUESIONER UTAMA"))

    return questions


def parse_to_json(docx_path: str, output_path: str | None = None) -> str:
    questions = parse_questionnaire(docx_path)
    q_ids = [q["id"] for q in questions if q["id"] and not q["is_header"]
             and q["route_type"] not in ("begin_group","end_group")]
    result = {
        "_meta": {
            "source":          Path(docx_path).name,
            "total_rows":      len(questions),
            "total_questions": len(q_ids),
            "question_ids":    q_ids,
        },
        "questions": questions,
    }
    js = json.dumps(result, ensure_ascii=False, indent=2)
    if output_path:
        Path(output_path).write_text(js, encoding="utf-8")
    return js


if __name__ == "__main__":
    fpath = sys.argv[1] if len(sys.argv) > 1 else None
    opath = sys.argv[2] if len(sys.argv) > 2 else None
    if not fpath:
        print("Usage: python docx_parser.py <input.docx> [output.json]")
        sys.exit(1)
    parse_to_json(fpath, opath)
    if opath:
        d = json.loads(Path(opath).read_text())
        m = d["_meta"]
        print(f"✓  {m['total_questions']} pertanyaan → {opath}")
        print(f"   First 15 IDs: {m['question_ids'][:15]}")